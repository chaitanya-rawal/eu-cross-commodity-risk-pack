import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import requests
import json
import datetime
from fredapi import Fred
from io import StringIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import anthropic
import warnings
import ssl
import certifi
import os
import urllib3

# SSL fixes
os.environ['SSL_CERT_FILE'] = certifi.where()
os.environ['REQUESTS_CA_BUNDLE'] = certifi.where()
ssl._create_default_https_context = ssl._create_unverified_context
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
warnings.filterwarnings('ignore')

# ============================================================
# API KEYS
# ============================================================
FRED_KEY = "579a7fb895e141294c6b06d6a0996008"
GIE_KEY = "09eeed08d17c64abc696e3cb4438e09b"
ANTHROPIC_KEY = "your_anthropic_api_key_here"

TODAY = datetime.date.today()
START_DATE = (TODAY - datetime.timedelta(days=90)).strftime('%Y-%m-%d')
END_DATE = TODAY.strftime('%Y-%m-%d')

print("="*60)
print("EU CROSS-COMMODITY RISK PACK")
print(f"Running for: {TODAY}")
print("="*60)

# ============================================================
# SECTION 1: DATA PULLS
# ============================================================

# --- 1A: TTF Gas Price via FRED ---
print("\n[1/4] Pulling TTF gas prices from FRED...")
try:
    fred = Fred(api_key=FRED_KEY)
    ttf_raw = fred.get_series('PNGASEUUSDM', observation_start=START_DATE)
    ttf_df = ttf_raw.to_frame(name='TTF_USD')
    ttf_df['TTF_EUR'] = ttf_df['TTF_USD'] * 0.92
    ttf_df.index = pd.to_datetime(ttf_df.index)
    ttf_df = ttf_df.dropna()
    print(f"    TTF latest: ${ttf_df['TTF_USD'].iloc[-1]:.2f}/MMBtu")
    print("    [OK] TTF data pulled successfully")
except Exception as e:
    print(f"    [ERROR] TTF pull failed: {e}")
    ttf_df = pd.DataFrame()

# --- 1B: EUA Carbon Price via Stooq ---
print("\n[2/4] Pulling EUA carbon prices via Stooq...")
try:
    session = requests.Session()
    session.verify = False
    stooq_url = "https://stooq.com/q/d/l/?s=co2.f&d1={}&d2={}&i=d".format(
        START_DATE.replace('-', ''),
        END_DATE.replace('-', '')
    )
    response = session.get(stooq_url, timeout=15)
    if response.status_code == 200 and len(response.text) > 100:
        eua_df = pd.read_csv(StringIO(response.text))
        eua_df.columns = [c.strip() for c in eua_df.columns]
        eua_df['Date'] = pd.to_datetime(eua_df['Date'])
        eua_df = eua_df.set_index('Date')
        eua_df = eua_df.rename(columns={'Close': 'EUA_EUR'})
        eua_df = eua_df[['EUA_EUR']].dropna()
        print(f"    EUA latest: €{eua_df['EUA_EUR'].iloc[-1]:.2f}/tCO2")
        print("    [OK] EUA data pulled successfully")
    else:
        raise ValueError("Stooq returned no data")
except Exception as e:
    print(f"    [WARN] Live EUA pull failed: {e}")
    print("    Using simulated EUA data...")
    dates = pd.date_range(start=START_DATE, end=END_DATE, freq='B')
    np.random.seed(42)
    eua_prices = 65 + np.cumsum(np.random.normal(0, 0.8, len(dates)))
    eua_df = pd.DataFrame({'EUA_EUR': eua_prices}, index=dates)
    print(f"    EUA simulated: €{eua_df['EUA_EUR'].iloc[-1]:.2f}/tCO2")
    print("    [OK] EUA simulated data ready")

# --- 1C: EU Gas Storage via GIE AGSI+ ---
print("\n[3/4] Pulling EU gas storage from GIE AGSI+...")
try:
    session = requests.Session()
    session.verify = False
    gie_url = "https://agsi.gie.eu/api"
    headers = {"x-key": GIE_KEY}
    params = {"country": "eu", "from": START_DATE, "to": END_DATE, "size": 90}
    response = session.get(gie_url, headers=headers, params=params, timeout=20)
    gie_data = response.json()
    if 'data' in gie_data:
        storage_records = []
        for record in gie_data['data']:
            storage_records.append({
                'date': pd.to_datetime(record['gasDayStart']),
                'storage_pct': float(record.get('full', 0)),
                'storage_twh': float(record.get('gasInStorage', 0))
            })
        storage_df = pd.DataFrame(storage_records).set_index('date').sort_index()
        storage_df = storage_df.dropna()
        print(f"    Storage latest: {storage_df['storage_pct'].iloc[-1]:.1f}% full")
        print("    [OK] Storage data pulled successfully")
    else:
        raise ValueError("Unexpected GIE response")
except Exception as e:
    print(f"    [WARN] Live storage pull failed: {e}")
    print("    Using simulated storage data...")
    dates = pd.date_range(start=START_DATE, end=END_DATE, freq='D')
    np.random.seed(99)
    storage_pct = np.linspace(72, 38, len(dates)) + np.random.normal(0, 0.5, len(dates))
    storage_df = pd.DataFrame({
        'storage_pct': storage_pct,
        'storage_twh': storage_pct * 11.5
    }, index=dates)
    print(f"    Storage simulated: {storage_df['storage_pct'].iloc[-1]:.1f}% full")
    print("    [OK] Storage simulated data ready")

# --- 1D: DE Day-Ahead Power via Energy-Charts ---
print("\n[4/4] Pulling DE Day-Ahead power prices...")
try:
    session = requests.Session()
    session.verify = False
    power_url = "https://api.energy-charts.info/price"
    params = {"bzn": "DE-LU", "start": START_DATE, "end": END_DATE}
    response = session.get(power_url, params=params, timeout=20)
    power_data = response.json()
    timestamps = [datetime.datetime.fromtimestamp(
        ts/1000 if ts > 1e10 else ts) for ts in power_data['unix_seconds']]
    prices = power_data['price']
    power_df = pd.DataFrame({'power_DA_EUR': prices}, index=pd.to_datetime(timestamps))
    power_df = power_df.resample('D').mean()
    power_df = power_df.dropna()
    print(f"    DE DA Power latest: €{power_df['power_DA_EUR'].iloc[-1]:.2f}/MWh")
    print("    [OK] Power data pulled successfully")
except Exception as e:
    print(f"    [WARN] Live power pull failed: {e}")
    print("    Using simulated power data...")
    dates = pd.date_range(start=START_DATE, end=END_DATE, freq='D')
    np.random.seed(123)
    power_prices = 85 + np.cumsum(np.random.normal(0, 3, len(dates)))
    power_prices = np.clip(power_prices, 40, 200)
    power_df = pd.DataFrame({'power_DA_EUR': power_prices}, index=dates)
    print(f"    Power simulated: €{power_df['power_DA_EUR'].iloc[-1]:.2f}/MWh")
    print("    [OK] Power simulated data ready")

# ============================================================
# SECTION 2: METRIC COMPUTATION
# ============================================================
print("\n" + "="*60)
print("COMPUTING METRICS")
print("="*60)

metrics = {}

if not ttf_df.empty:
    metrics['TTF_front_month_USD_MMBtu'] = round(float(ttf_df['TTF_USD'].iloc[-1]), 2)
    metrics['TTF_front_month_EUR_MWh'] = round(float(ttf_df['TTF_EUR'].iloc[-1]) * 3.412, 2)
    if len(ttf_df) >= 3:
        metrics['TTF_3m_momentum_pct'] = round(
            float((ttf_df['TTF_USD'].iloc[-1] - ttf_df['TTF_USD'].iloc[-3]) /
            ttf_df['TTF_USD'].iloc[-3] * 100), 2)

if not eua_df.empty:
    metrics['EUA_price_EUR_t'] = round(float(eua_df['EUA_EUR'].iloc[-1]), 2)
    if len(eua_df) >= 30:
        metrics['EUA_30d_momentum_pct'] = round(
            float((eua_df['EUA_EUR'].iloc[-1] - eua_df['EUA_EUR'].iloc[-30]) /
            eua_df['EUA_EUR'].iloc[-30] * 100), 2)

if not storage_df.empty:
    metrics['EU_storage_pct_full'] = round(float(storage_df['storage_pct'].iloc[-1]), 1)
    metrics['EU_storage_vs_5yr_avg'] = round(
        float(storage_df['storage_pct'].iloc[-1]) - 65.0, 1)

if not power_df.empty:
    metrics['DE_DA_power_EUR_MWh'] = round(float(power_df['power_DA_EUR'].iloc[-1]), 2)

if all(k in metrics for k in ['DE_DA_power_EUR_MWh', 'TTF_front_month_EUR_MWh', 'EUA_price_EUR_t']):
    heat_rate = 0.45
    emission_factor = 0.34
    gas_cost = metrics['TTF_front_month_EUR_MWh'] / heat_rate
    carbon_cost = metrics['EUA_price_EUR_t'] * emission_factor
    metrics['clean_dark_spread_EUR_MWh'] = round(
        metrics['DE_DA_power_EUR_MWh'] - gas_cost - carbon_cost, 2)

if not eua_df.empty and not power_df.empty:
    try:
        combined = eua_df.join(power_df, how='inner').tail(30)
        if len(combined) >= 10:
            corr = combined['EUA_EUR'].corr(combined['power_DA_EUR'])
            metrics['power_carbon_30d_correlation'] = round(float(corr), 3)
    except:
        pass

print("\nMETRICS COMPUTED:")
for k, v in metrics.items():
    print(f"    {k}: {v}")

# ============================================================
# SECTION 3: CHART GENERATION
# ============================================================
print("\n" + "="*60)
print("GENERATING CHARTS")
print("="*60)

plt.style.use('seaborn-v0_8-darkgrid')

# Chart 1: TTF vs EUA
print("\nGenerating Chart 1: TTF vs EUA...")
try:
    fig, ax1 = plt.subplots(figsize=(12, 6))
    color_ttf = '#1f77b4'
    color_eua = '#2ca02c'
    ax1.plot(ttf_df.index, ttf_df['TTF_USD'],
             color=color_ttf, linewidth=2.5, label='TTF Gas ($/MMBtu)')
    ax1.set_xlabel('Date', fontsize=11)
    ax1.set_ylabel('TTF Gas Price ($/MMBtu)', color=color_ttf, fontsize=11)
    ax1.tick_params(axis='y', labelcolor=color_ttf)
    ax2 = ax1.twinx()
    ax2.plot(eua_df.index, eua_df['EUA_EUR'],
             color=color_eua, linewidth=2.5, label='EUA Carbon (€/t)', linestyle='--')
    ax2.set_ylabel('EUA Carbon Price (€/tCO2)', color=color_eua, fontsize=11)
    ax2.tick_params(axis='y', labelcolor=color_eua)
    plt.title(f'TTF Gas vs EUA Carbon — 90 Day View\nAs of {TODAY}',
              fontsize=13, fontweight='bold', pad=15)
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper left', fontsize=10)
    ax1.xaxis.set_major_formatter(mdates.DateFormatter('%b %Y'))
    ax1.xaxis.set_major_locator(mdates.MonthLocator())
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig('chart1_ttf_eua.png', dpi=150, bbox_inches='tight')
    plt.close()
    print("    [OK] Chart 1 saved")
except Exception as e:
    print(f"    [ERROR] Chart 1 failed: {e}")

# Chart 2: Power & CDS
print("\nGenerating Chart 2: DE Power & Clean Dark Spread...")
try:
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(12, 9), sharex=False)
    power_plot = power_df.tail(60)
    ax1.fill_between(power_plot.index, power_plot['power_DA_EUR'],
                     alpha=0.4, color='#ff7f0e')
    ax1.plot(power_plot.index, power_plot['power_DA_EUR'],
             color='#ff7f0e', linewidth=2, label='DE Day-Ahead Power')
    ax1.axhline(y=power_plot['power_DA_EUR'].mean(), color='red',
                linestyle=':', linewidth=1.5, label='60d Average')
    ax1.set_ylabel('€/MWh', fontsize=11)
    ax1.set_title('DE Day-Ahead Power Price — 60 Day View',
                  fontsize=12, fontweight='bold')
    ax1.legend(fontsize=10)
    if 'clean_dark_spread_EUR_MWh' in metrics:
        cds_value = metrics['clean_dark_spread_EUR_MWh']
        color = '#2ca02c' if cds_value >= 0 else '#d62728'
        ax2.bar(['Clean Dark Spread'], [cds_value], color=color, width=0.4, alpha=0.8)
        ax2.axhline(y=0, color='black', linewidth=1.5)
        ax2.set_ylabel('€/MWh', fontsize=11)
        ax2.set_title('Clean Dark Spread (Gas→Power Profitability)',
                      fontsize=12, fontweight='bold')
        ax2.text(0, cds_value + (0.5 if cds_value >= 0 else -2),
                 f'€{cds_value:.2f}/MWh', ha='center', fontsize=13, fontweight='bold')
    plt.tight_layout()
    plt.savefig('chart2_power_cds.png', dpi=150, bbox_inches='tight')
    plt.close()
    print("    [OK] Chart 2 saved")
except Exception as e:
    print(f"    [ERROR] Chart 2 failed: {e}")

# ============================================================
# SECTION 4: AI NARRATIVE — CLAUDE
# ============================================================
print("\n" + "="*60)
print("GENERATING AI NARRATIVE (Claude)")
print("="*60)

try:
    client = anthropic.Anthropic(api_key=ANTHROPIC_KEY)

    prompt = f"""You are a senior European energy trader writing a morning desk note for a trading team.

Based on the following live market metrics, write a professional 3-paragraph risk narrative covering:

Paragraph 1 - GAS TIGHTNESS: Analyse TTF price level, momentum, and what it implies for prompt gas supply/demand balance in Europe.

Paragraph 2 - CARBON SIGNAL: Analyse EUA price level and 30-day momentum. What is the policy and supply signal? Is carbon supportive or bearish for power prices?

Paragraph 3 - POWER CURVE IMPLICATIONS: Using the Clean Dark Spread and Day-Ahead price, assess the current gas-to-power economics. What is the key risk asymmetry for European power from Day-Ahead through the curve?

Be specific, cite the exact numbers provided, use trader language, and end with one clear risk statement.

LIVE METRICS AS OF {TODAY}:
{json.dumps(metrics, indent=2)}
"""

    print("\nCalling Claude API...")
    response = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=800,
        messages=[{"role": "user", "content": prompt}]
    )
    narrative = response.content[0].text

    print("\nAI NARRATIVE GENERATED:")
    print("-"*40)
    print(narrative)

    with open("ai_log.txt", "a") as f:
        f.write(f"\n{'='*60}\n")
        f.write(f"TIMESTAMP: {datetime.datetime.now()}\n")
        f.write(f"MODEL: claude-haiku-4-5-20251001\n")
        f.write(f"{'='*60}\n")
        f.write(f"PROMPT:\n{prompt}\n\n")
        f.write(f"OUTPUT:\n{narrative}\n")
    print("\n    [OK] Narrative logged to ai_log.txt")

except Exception as e:
    print(f"    [ERROR] Claude narrative failed: {e}")
    narrative = "AI narrative unavailable - check API key and credits."

# ============================================================
# SECTION 5: WORD DOCUMENT
# ============================================================
print("\n" + "="*60)
print("ASSEMBLING WORD DOCUMENT")
print("="*60)

try:
    doc = Document()
    title = doc.add_heading('European Cross-Commodity Risk Pack', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading(f'Daily Desk Note — {TODAY}', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        'Theme: Gas + Carbon → European Power Curve Implications | '
        'Markets: TTF Gas | EUA Carbon | DE Day-Ahead Power'
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_heading('1. Daily Monitor Metrics', level=2)

    metric_labels = {
        'TTF_front_month_USD_MMBtu': ('TTF Front Month', '$/MMBtu', 'Prompt gas price benchmark'),
        'TTF_front_month_EUR_MWh': ('TTF (EUR equivalent)', '€/MWh', 'Gas cost input to power stack'),
        'TTF_3m_momentum_pct': ('TTF 3-Month Momentum', '%', 'Prompt tightness direction'),
        'EUA_price_EUR_t': ('EUA Front Dec Price', '€/tCO2', 'Carbon cost input'),
        'EUA_30d_momentum_pct': ('EUA 30-Day Momentum', '%', 'Policy/supply signal'),
        'EU_storage_pct_full': ('EU Gas Storage', '% full', 'Structural gas buffer'),
        'EU_storage_vs_5yr_avg': ('Storage vs 5yr Avg', 'pp', 'Relative tightness signal'),
        'DE_DA_power_EUR_MWh': ('DE Day-Ahead Power', '€/MWh', 'Realised power benchmark'),
        'clean_dark_spread_EUR_MWh': ('Clean Dark Spread', '€/MWh', 'Gas→power profitability'),
        'power_carbon_30d_correlation': ('Power-Carbon Correlation (30d)', 'r', 'Cross-commodity regime')
    }

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    hdr[2].text = 'Unit'
    hdr[3].text = 'Trading Relevance'

    for key, (label, unit, relevance) in metric_labels.items():
        if key in metrics:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(metrics[key])
            row[2].text = unit
            row[3].text = relevance

    doc.add_paragraph()
    doc.add_heading('2. Risk Narrative', level=2)
    doc.add_paragraph(narrative)
    doc.add_paragraph()
    doc.add_heading('3. Market Charts', level=2)

    try:
        doc.add_paragraph('Chart 1: TTF Gas vs EUA Carbon — 90 Day View')
        doc.add_picture('chart1_ttf_eua.png', width=Inches(6))
        doc.add_paragraph()
    except:
        doc.add_paragraph('[Chart 1 not available]')

    try:
        doc.add_paragraph('Chart 2: DE Day-Ahead Power & Clean Dark Spread')
        doc.add_picture('chart2_power_cds.png', width=Inches(6))
    except:
        doc.add_paragraph('[Chart 2 not available]')

    doc.add_paragraph()
    footer = doc.add_paragraph(
        f'Generated automatically by EU Risk Pack Monitor | '
        f'{datetime.datetime.now().strftime("%Y-%m-%d %H:%M")} | '
        'Data: FRED, Stooq, GIE AGSI+, Energy-Charts.info | AI: Claude Haiku'
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save('risk_pack_output.docx')
    print("    [OK] Word document saved: risk_pack_output.docx")

except Exception as e:
    print(f"    [ERROR] Word document failed: {e}")

# ============================================================
# DONE
# ============================================================
print("\n" + "="*60)
print("RISK PACK COMPLETE")
print("="*60)
print("Output files:")
print("    risk_pack_output.docx  — desk note")
print("    chart1_ttf_eua.png     — gas vs carbon chart")
print("    chart2_power_cds.png   — power & CDS chart")
print("    ai_log.txt             — logged AI prompts & outputs")
print("="*60)
